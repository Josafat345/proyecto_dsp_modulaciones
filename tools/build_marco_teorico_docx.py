from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "marco_teorico_y_justificacion.md"
OUTPUT = ROOT / "outputs" / "word" / "Marco_Teorico_y_Justificacion_DSP_IA.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "F4F6F9"
BORDER = "CBD5E1"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph_with_inline_markdown(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text)
    return p


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, color=INK)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_BLUE
        else:
            run = paragraph.add_run(part)
            set_run_font(run, color=INK)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    add_inline_runs(p2, body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("Documento", "Marco teorico y justificacion"),
        ("Proyecto", "Reconocimiento inteligente de modulaciones digitales con DSP e IA"),
        ("Enfoque", "Clasificacion automatica de modulaciones, rasgos DSP, MLP ligero y pruebas de robustez"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060], 120)
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), TABLE_FILL)
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10.5, color=INK, bold=(cell == table.cell(i, 0)))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(15, 23, 42)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Calibri"
        callout.font.size = Pt(10.5)
        callout.font.color.rgb = DARK_BLUE
        callout.paragraph_format.space_after = Pt(4)
        callout.paragraph_format.line_spacing = 1.208


def add_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = "Proyecto DSP + IA | Marco teorico y justificacion"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Marco teorico y justificacion | Proyecto DSP + IA")
    for run in p.runs:
        set_run_font(run, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph("DOCUMENTO DE APOYO PARA PROPUESTA DE INVESTIGACION")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in kicker.runs:
        set_run_font(run, size=10.5, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.add_run("Marco teorico y justificacion del proyecto")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Evaluacion de robustez y explicabilidad de un clasificador inteligente de modulaciones digitales basado en rasgos DSP bajo condiciones de canal no ideales")
    set_run_font(run, size=13, italic=True, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(24)

    add_metadata_table(doc)

    for _ in range(2):
        doc.add_paragraph()
    add_callout(
        doc,
        "Resumen ejecutivo",
        "Este documento explica los fundamentos teoricos del proyecto, justifica su relevancia academica y tecnica, y describe por que los experimentos propuestos permiten defender la investigacion ante un asesor.",
    )
    doc.add_page_break()


def add_static_contents(doc: Document, headings: list[tuple[int, str]]) -> None:
    doc.add_heading("Contenido", level=1)
    for level, title in headings:
        if level == 2:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, title)
        elif level == 3 and title.startswith(("3.", "5.", "8.")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.55)
            add_inline_runs(p, title)
    note = doc.add_paragraph()
    note.style = "Callout"
    note.add_run("Nota: ").bold = True
    note.add_run("este contenido es una guia estatica de navegacion; los titulos reales se encuentran en las paginas siguientes.")
    doc.add_page_break()


def collect_headings(markdown: str) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        if line.startswith("## "):
            headings.append((2, line[3:].strip()))
        elif line.startswith("### "):
            headings.append((3, line[4:].strip()))
    return headings


def markdown_to_docx(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            # The cover already carries the document title.
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("> "):
            quote = doc.add_paragraph(style="Intense Quote")
            add_inline_runs(quote, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        add_paragraph_with_inline_markdown(doc, stripped)
        i += 1


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)
    add_static_contents(doc, collect_headings(markdown))
    markdown_to_docx(doc, markdown)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
